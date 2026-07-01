from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn


SRC = Path("/Users/liuziyu/Nz/Resume_2026_Ziyu_Liu_template_style.docx")
OUT = Path("/Users/liuziyu/Nz/Resume_2026_Ziyu_Liu_IAG_Digital_Content_Specialist.docx")


def font_from(run):
    return {
        "name": run.font.name,
        "size": run.font.size,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "color": run.font.color.rgb,
    }


def apply_font(run, fmt):
    if fmt["name"]:
        run.font.name = fmt["name"]
        run._element.rPr.rFonts.set(qn("w:ascii"), fmt["name"])
        run._element.rPr.rFonts.set(qn("w:hAnsi"), fmt["name"])
    run.font.size = fmt["size"]
    run.bold = fmt["bold"]
    run.italic = fmt["italic"]
    run.underline = fmt["underline"]
    if fmt["color"]:
        run.font.color.rgb = fmt["color"]


def replace_text(paragraph, text, fmt_source=None):
    fmt = font_from(fmt_source or paragraph.runs[0]) if paragraph.runs else None
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    if fmt:
        apply_font(run, fmt)
    return run


def replace_skill(paragraph, label, value):
    label_fmt = font_from(paragraph.runs[0]) if paragraph.runs else None
    value_fmt = font_from(paragraph.runs[-1]) if paragraph.runs else label_fmt
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    r = paragraph.add_run(label)
    if label_fmt:
        apply_font(r, label_fmt)
    r = paragraph.add_run(": ")
    if label_fmt:
        apply_font(r, label_fmt)
    r = paragraph.add_run(value)
    if value_fmt:
        apply_font(r, value_fmt)


def main():
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)

    # Profile: keep the same paragraph but position the CV toward the IAG role.
    replace_text(
        doc.paragraphs[2],
        "Digital content and technology professional with hands-on experience maintaining www.nzddvisa.com, managing website content, SEO structure, digital publishing workflows, engagement analytics, and cross-platform content operations. Master of IT at the University of Auckland, with a strong foundation in web systems, data analysis, cloud tools, and technical problem-solving. Able to translate business and customer needs into clear, reliable digital experiences, combining content quality, stakeholder collaboration, attention to detail, and performance-informed improvement.",
    )

    # Skills: preserve section layout while surfacing Digital Content Specialist keywords.
    replace_skill(doc.paragraphs[4], "Digital Content & Web", "Website content maintenance, digital publishing workflows, customer-focused web pages, content quality control, SEO structure")
    replace_skill(doc.paragraphs[5], "Technical Skills", "HTML, CSS, JavaScript, React/MERN, Python, SQL, Git, basic CMS-style content operations")
    replace_skill(doc.paragraphs[6], "SEO & Analytics", "SEO fundamentals, engagement analytics, funnel metrics, performance monitoring, data-informed content improvement")
    replace_skill(doc.paragraphs[7], "Content Production", "Photography, video editing, Photoshop, CapCut, RedNote, TikTok, multimedia campaign assets")
    replace_skill(doc.paragraphs[8], "Cloud & Tooling", "AWS Lambda, Amazon EC2, DynamoDB, Bedrock, S3, Git, ELK, Kibana, JMeter, Docker")

    # Current DD Immigration role: tailor toward IAG Digital Content Specialist responsibilities.
    replace_text(
        doc.paragraphs[11],
        "Digital Marketing Technology Specialist                                                     \t\t        Auckland, Jan 2026 - Present",
    )
    replace_text(doc.paragraphs[12], "DD Immigration | www.nzddvisa.com")
    dd_points = [
        "Maintain and continuously improve www.nzddvisa.com, managing website content updates, page structure, SEO foundations, performance monitoring, and day-to-day site reliability.",
        "Produce, update, and publish clear digital content across the company website and social channels, ensuring accuracy, consistency, audience relevance, and a professional customer experience.",
        "Build practical content workflows across RedNote and TikTok, coordinating production, review, publishing, asset management, and engagement analytics.",
        "Use SEO, engagement data, funnel metrics, and website performance signals to improve discoverability, content effectiveness, and digital user journeys.",
        "Collaborate with consultants and business stakeholders to turn immigration service information into clear, compliant, customer-friendly content, supported by documentation, content library management, data backup, and pre-publication quality checks.",
    ]
    for idx, text in enumerate(dd_points, start=13):
        replace_text(doc.paragraphs[idx], text)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
