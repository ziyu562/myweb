from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/liuziyu/Nz/Tracy Kong CV - IAG Digital Content Specialist.docx")


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, before=0, after=5, line=1.08, left_indent=None, hanging=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging is not None:
        pf.first_line_indent = hanging


def add_bottom_border(paragraph, color="BFBFBF", size="6", space="3"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=10, after=4, line=1.0)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color="000000")
    add_bottom_border(p, color="D9D9D9", size="4", space="2")
    return p


def add_body(doc, text, after=5):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.08)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    set_para(p, after=2.5, line=1.05, left_indent=Cm(0.55), hanging=Cm(-0.22))
    r = p.add_run("• ")
    set_font(r, size=10.2)
    r = p.add_run(text)
    set_font(r, size=10.2)
    return p


def add_role(doc, title, dates, organisation=None, location=None, summary=None, bullets=None):
    p = doc.add_paragraph()
    set_para(p, before=5, after=0, line=1.0)
    left = p.add_run(title)
    set_font(left, size=10.8, bold=True)
    tab = p.add_run("\t")
    set_font(tab, size=10.8)
    right = p.add_run(dates)
    set_font(right, size=10.2)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0))

    if organisation:
        p2 = doc.add_paragraph()
        set_para(p2, after=2, line=1.0)
        text = organisation if not location else f"{organisation}, {location}"
        r = p2.add_run(text)
        set_font(r, size=10.2, italic=True, color="404040")
    if summary:
        add_body(doc, f"Role Summary: {summary}", after=3)
    for item in bullets or []:
        add_bullet(doc, item)


def add_skill_line(doc, label, value):
    p = doc.add_paragraph()
    set_para(p, after=2, line=1.05)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.3, bold=True)
    r = p.add_run(value)
    set_font(r, size=10.3)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(title, after=2, line=1.0)
    r = title.add_run("Tracy Kong")
    set_font(r, size=20, bold=True)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(contact, after=8, line=1.0)
    r = contact.add_run("Mobile: 029 0250 2691  |  Email: jkon419@aucklanduni.ac.nz  |  Location: Auckland, New Zealand")
    set_font(r, size=9.8, color="404040")

    add_heading(doc, "Profile")
    add_body(
        doc,
        "Digital content and media professional with 7+ years of experience in content planning, scriptwriting, video production, editing, and multi-platform publishing. Experienced in producing clear, engaging content for television, short-form video, news, and digital channels including YouTube, WeChat Video, and Douyin. Skilled at turning business, audience, and stakeholder needs into practical content plans, managing production and review workflows, and delivering accurate, high-quality content under deadline pressure. Strong fit for digital content roles requiring attention to detail, customer-focused communication, cross-functional collaboration, content optimisation, and continuous improvement of digital journeys.",
        after=6,
    )

    add_heading(doc, "Skills")
    add_skill_line(doc, "Digital Content", "Website and digital platform content, content planning, publishing workflows, copy/scriptwriting, storytelling, editorial consistency")
    add_skill_line(doc, "Customer Experience", "Audience-focused communication, clear information design, content quality control, brand and message alignment")
    add_skill_line(doc, "CMS & Publishing", "Digital content upload and maintenance, platform publishing, YouTube, WeChat Video, Douyin; confident learning CMS platforms such as AEM")
    add_skill_line(doc, "SEO & Analytics", "SEO fundamentals, discoverability-focused content, audience feedback review, engagement and performance-informed improvement")
    add_skill_line(doc, "Production Tools", "Final Cut Pro, Adobe Premiere Pro, CapCut, Vmix, DSLR cameras, gimbals/stabilizers, Adobe After Effects (basic), Adobe Photoshop (basic)")
    add_skill_line(doc, "Professional Skills", "Stakeholder communication, project coordination, approval workflows, attention to detail, deadline management, teamwork")
    add_skill_line(doc, "Languages", "Mandarin, English, Korean")

    add_heading(doc, "WORK EXPERIENCE")
    add_role(
        doc,
        "Post-production and Content Creator",
        "Aug 2025 - Present",
        "Channel 33",
        "Auckland, New Zealand",
        "Plan, produce, edit, and publish television, news, and short-form content for a Chinese-language media organisation in New Zealand, supporting both broadcast and digital distribution.",
        [
            "Produce and publish digital content across YouTube, WeChat Video, and Douyin, adapting content format and messaging for each platform.",
            "Plan, shoot, interview, and edit news and feature content, ensuring accuracy, clarity, and consistent production quality.",
            "Serve as live program director using Vmix, coordinating multi-camera production and real-time broadcast delivery.",
            "Collaborate with production teams on concept development, scriptwriting, filming, editing, review, and final delivery.",
            "Maintain high attention to detail across subtitles, visuals, timing, audio, and platform-ready outputs before publication.",
        ],
    )
    add_role(
        doc,
        "Sales Associate (Part-time)",
        "Oct 2024 - Jun 2025",
        "Danik Bathroom",
        "Auckland, New Zealand",
        "Supported customer communication, product explanation, and promotional activity in a retail environment.",
        [
            "Explained product features and recommendations clearly to customers, translating needs into practical options.",
            "Maintained customer relationships through order follow-up and after-sales service, supporting satisfaction and trust.",
            "Assisted with product display, store presentation, and promotional execution to improve customer experience.",
        ],
    )
    add_role(
        doc,
        "Short Video Director",
        "Nov 2021 - Oct 2023",
        "Tomorrow Advancing Life",
        "Beijing, China",
        "Directed short-form video projects from concept to final delivery, with responsibility for audience-focused ideas, scripts, production planning, filming, and post-production.",
        [
            "Developed creative concepts and scripts tailored to target audiences, project goals, and brand messaging.",
            "Managed preparation workflows including location scouting, talent casting, budgeting, scheduling, and production coordination.",
            "Directed actors, camera operators, and crew during filming, maintaining time, quality, and creative standards.",
            "Worked with editors, sound designers, and visual effects artists to deliver polished final content.",
            "Liaised with clients and stakeholders, presenting concepts, progress updates, and final deliverables clearly.",
        ],
    )
    add_role(
        doc,
        "Variety Show Director",
        "Sep 2019 - Oct 2021",
        "Beijing Yuanchun Media",
        "Beijing, China",
        "Led creative direction and production workflows for variety show content.",
        [
            "Developed concepts, scripts, and show formats in collaboration with writers, producers, marketing, and distribution teams.",
            "Directed on-site filming and coordinated production staff, technical crew, and performers under tight deadlines.",
            "Supervised editing, sound, and visual effects to ensure final content met creative and technical standards.",
            "Reviewed audience feedback and ratings to refine future content and improve audience appeal.",
        ],
    )
    add_role(
        doc,
        "Variety Show Director",
        "Jul 2017 - Aug 2019",
        "Youku",
        "Beijing, China",
        "Supported full-cycle variety show planning and production for a major digital video platform.",
        [
            "Conducted topic research and pre-show planning aligned with programme positioning and audience expectations.",
            "Collected, organised, and maintained programme materials to support consistent style and smooth production.",
            "Led scriptwriting coordination and post-production communication to maintain quality across the production workflow.",
            "Monitored production progress to ensure final programmes met delivery standards.",
        ],
    )

    add_heading(doc, "Selected Content Projects")
    add_role(
        doc,
        "Digital Video and Broadcast Content Portfolio",
        "2017 - Present",
        summary="Contributed to 20+ variety shows, news items, short-form videos, and digital content projects across China and New Zealand.",
        bullets=[
            "Delivered content across the full lifecycle: idea development, scripting, filming, editing, review, publishing, and improvement.",
            "Balanced creative quality with operational requirements including scheduling, approvals, platform requirements, and audience expectations.",
            "Built strong transferable experience for website and digital content work: clarity, accuracy, consistency, stakeholder review, and customer-focused presentation.",
        ],
    )

    add_heading(doc, "EDUCATION")
    add_role(
        doc,
        "Master of Media and Screen Studies",
        "",
        "The University of Auckland",
        "Auckland, New Zealand",
    )
    add_role(
        doc,
        "Bachelor of Radio and Television Directing",
        "",
        "Jilin College of the Arts",
        "Jilin, China",
    )

    add_heading(doc, "REFEREE")
    add_body(doc, "Referees are available upon request.", after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
