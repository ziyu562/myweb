from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCX = Path("/Users/liuziyu/Nz/Tracy Kong CV - IAG Digital Content Specialist.docx")
AVATAR = Path("/Users/liuziyu/myweb/.work/cv_tracy_iag/assets/tracy_avatar_hi.png")
DAILY = Path("/Users/liuziyu/myweb/.work/cv_tracy_iag/assets/tracy_my_daily_life.png")
BACKUP = DOCX.with_name(f"{DOCX.stem}.backup-visuals-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx")


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.add_paragraph()._p.__class__(new_p, paragraph._parent)


def add_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.paragraphs[-1]._parent._body._body  # not used


def paragraph_after(doc, paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def set_run_font(run, size=12, bold=False, color="404040"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_section_heading(paragraph, text):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_run_font(run, size=16, bold=True, color="404040")


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    # Add the small female avatar near the top, similar to Hughie's CV.
    contact = doc.paragraphs[1]
    avatar_p = paragraph_after(doc, contact)
    avatar_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    avatar_p.paragraph_format.space_before = Pt(0)
    avatar_p.paragraph_format.space_after = Pt(2)
    avatar_p.add_run().add_picture(str(AVATAR), width=Inches(1.0))

    # Add a My Daily Life visual after education and before referee.
    referee_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "REFEREE")
    before_ref = doc.paragraphs[referee_idx - 1]
    spacer = paragraph_after(doc, before_ref)
    spacer.paragraph_format.space_after = Pt(2)

    heading = paragraph_after(doc, spacer)
    add_section_heading(heading, "My Daily Life")

    image_p = paragraph_after(doc, heading)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_before = Pt(2)
    image_p.paragraph_format.space_after = Pt(8)
    image_p.add_run().add_picture(str(DAILY), width=Inches(6.35))

    doc.save(DOCX)
    print(BACKUP)


if __name__ == "__main__":
    main()
