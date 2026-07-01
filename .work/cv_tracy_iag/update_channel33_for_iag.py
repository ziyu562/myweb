from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn


DOCX = Path("/Users/liuziyu/Nz/Tracy Kong CV - IAG Digital Content Specialist.docx")
BACKUP = DOCX.with_name(
    f"{DOCX.stem}.backup-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"
)


def run_format(run):
    return {
        "name": run.font.name,
        "size": run.font.size,
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "color": run.font.color.rgb,
    }


def apply_format(run, fmt):
    if fmt["name"]:
        run.font.name = fmt["name"]
        run._element.rPr.rFonts.set(qn("w:ascii"), fmt["name"])
        run._element.rPr.rFonts.set(qn("w:hAnsi"), fmt["name"])
    run.font.size = fmt["size"]
    run.bold = fmt["bold"]
    run.italic = fmt["italic"]
    run.underline = fmt["underline"]
    if fmt["color"]:
        run.font.color.rgb = fmt["color"]


def replace_paragraph(paragraph, text):
    fmt = run_format(paragraph.runs[-1] if paragraph.runs else paragraph.add_run())
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    apply_format(run, fmt)


def main():
    shutil.copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    replacements = {
        15: "Role Summary: Produce, optimise, and publish broadcast and digital content for Channel 33, supporting website-ready stories, social video, news packages, and multi-platform audience engagement. The role combines content planning, editorial accuracy, production workflows, performance awareness, and quality control before publication.",
        16: "• Produce and publish digital content across YouTube, WeChat Video, and Douyin, adapting titles, captions, descriptions, thumbnails, and formats to improve discoverability and engagement.",
        17: "• Apply SEO-aware content practices, including clear headlines, audience-focused descriptions, keyword-conscious wording, and consistent metadata for platform search and viewer navigation.",
        18: "• Review audience feedback, platform performance signals, and content engagement patterns to refine topics, packaging, timing, and future digital content decisions.",
        19: "• Coordinate content workflows from planning, shooting, interviewing, editing, review, approval, and final publishing, maintaining accuracy and delivery standards under deadline pressure.",
        20: "• Collaborate with producers, reporters, presenters, and technical teams to turn complex news or programme ideas into clear, customer-friendly digital content with strong attention to subtitles, visuals, audio, and brand consistency.",
    }

    for idx, text in replacements.items():
        replace_paragraph(doc.paragraphs[idx], text)

    doc.save(DOCX)
    print(BACKUP)


if __name__ == "__main__":
    main()
